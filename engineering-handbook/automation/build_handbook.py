"""Build source Markdown into an authoritative DOCX and derived PDF."""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Iterable

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader

from document_model import Node


SOFFICE = Path("/Users/aryansingh/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/override/soffice")
PDFTOPPM = shutil.which("pdftoppm") or "/opt/homebrew/bin/pdftoppm"
STYLE_PATH = Path(__file__).parents[1] / "build/styles/publication.yaml"
LINK = re.compile(r"\[[^\]]+\]\(([^)#]+)")
TASK = re.compile(r"^\s*[-*]\s+\[([ xX])\]\s+(.+)$")
INLINE_LINK = re.compile(r"\[([^\]]+)\]\([^)]*\)")


def _read_yaml(path: Path) -> dict:
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    return data if isinstance(data, dict) else {}


def manifest_paths(root: Path) -> list[Path]:
    return [Path(value) for value in LINK.findall((root / "SUMMARY.md").read_text(encoding="utf-8"))]


def manifest_headings(root: Path, paths: Iterable[Path] | None = None) -> list[Node]:
    """Return one navigation heading per manifest document, in publication order."""
    headings: list[Node] = []
    for relative in list(paths or manifest_paths(root)):
        nodes = parse_document((root / relative).read_text(encoding="utf-8"))
        heading = next((node for node in nodes if node.kind == "heading"), None)
        if heading is not None:
            headings.append(heading)
    return headings


def strip_front_matter(content: str) -> str:
    if content.startswith("---\n"):
        return content.split("---\n", 2)[2]
    return content


def parse_document(content: str) -> list[Node]:
    """Parse the supported canonical Markdown subset into semantic nodes."""
    lines, nodes, index = strip_front_matter(content).splitlines(), [], 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("```"):
            language, code = line[3:].strip(), []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code.append(lines[index])
                index += 1
            nodes.append(Node("mermaid" if language == "mermaid" else "code", "\n".join(code), language=language))
            index += 1
            continue
        if match := re.match(r"^(#{1,6})\s+(.+)$", line):
            nodes.append(Node("heading", match.group(2).strip(), level=len(match.group(1))))
        elif match := TASK.match(line):
            nodes.append(Node("task", match.group(2), checked=match.group(1).lower() == "x"))
        elif re.match(r"^\s*[-*]\s+", line):
            nodes.append(Node("bullet", re.sub(r"^\s*[-*]\s+", "", line)))
        elif re.match(r"^\s*\d+\.\s+", line):
            nodes.append(Node("ordered", re.sub(r"^\s*\d+\.\s+", "", line)))
        elif line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-{3,}", lines[index + 1]):
            headers = [cell.strip() for cell in line.strip("|").split("|")]
            rows, index = [headers], index + 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip("|").split("|")])
                index += 1
            nodes.append(Node("table", rows=rows))
            continue
        elif line.startswith("> "):
            nodes.append(Node("callout", line[2:].strip()))
        elif line == ":::pagebreak":
            nodes.append(Node("pagebreak"))
        elif line == ":::volumebreak":
            nodes.append(Node("volumebreak"))
        else:
            nodes.append(Node("paragraph", line.strip()))
        index += 1
    return nodes


def build_markdown(root: Path, output: Path, paths: Iterable[Path] | None = None) -> Path:
    root = root.resolve()
    title = _read_yaml(root / "metadata.yaml").get("handbook", {}).get("title", "Enterprise Engineering Manual")
    selected = list(paths or manifest_paths(root))
    bodies = [f"# {title}"]
    for relative in selected:
        bodies.append(strip_front_matter((root / relative).read_text(encoding="utf-8")).strip())
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n\n".join(bodies).rstrip() + "\n", encoding="utf-8")
    return output


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _anchor(text: str) -> str:
    return "h-" + re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")[:36]


def _bookmark(paragraph, anchor: str, bookmark_id: int) -> None:
    start, end = OxmlElement("w:bookmarkStart"), OxmlElement("w:bookmarkEnd")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), anchor)
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_link(paragraph, text: str, anchor: str) -> None:
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4D78")
    properties.append(color)
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    link.append(run)
    paragraph._p.append(link)


def _configure_styles(document: Document) -> None:
    tokens = _read_yaml(STYLE_PATH)
    typeface = tokens["typography"]["body_font"]
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name, normal.font.size = typeface, Pt(tokens["typography"]["body_size_pt"])
    normal.paragraph_format.space_after = Pt(tokens["paragraph"]["after_pt"])
    normal.paragraph_format.line_spacing = tokens["paragraph"]["line_spacing"]
    for number, size, color in ((1, 16, "2E74B5"), (2, 13, "2E74B5"), (3, 12, "1F4D78")):
        style = styles[f"Heading {number}"]
        style.font.name, style.font.size, style.font.bold = typeface, Pt(size), True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(14 if number == 1 else 10), Pt(6)
    if "Handbook Code" not in styles:
        code = styles.add_style("Handbook Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name, code.font.size = tokens["typography"]["code_font"], Pt(8)
        code.paragraph_format.space_after = Pt(6)


def _set_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.8)
    section.header_distance = section.footer_distance = Inches(0.35)
    header = section.header.paragraphs[0]
    header.text = "Enterprise Engineering Manual"
    header.style = document.styles["Normal"]
    header.runs[0].font.size, header.runs[0].font.color.rgb = Pt(8), RGBColor(100, 100, 100)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Internal engineering reference | ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def _append_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    for row_index, values in enumerate(rows):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            if row_index == 0:
                _set_cell_shading(cells[index], "E8EEF5")
    document.add_paragraph()


def _append_node(document: Document, node: Node, page_break_before: bool = False, bookmark_id: int = 0) -> None:
    if node.kind == "heading":
        paragraph = document.add_paragraph(node.text, style=f"Heading {min(node.level, 3)}")
        paragraph.paragraph_format.page_break_before = page_break_before
        _bookmark(paragraph, _anchor(node.text), bookmark_id)
    elif node.kind == "paragraph":
        document.add_paragraph(INLINE_LINK.sub(r"\1", node.text))
    elif node.kind == "bullet":
        document.add_paragraph(node.text, style="List Bullet")
    elif node.kind == "ordered":
        document.add_paragraph(node.text, style="List Number")
    elif node.kind == "task":
        document.add_paragraph(("☒ " if node.checked else "☐ ") + node.text, style="List Bullet")
    elif node.kind == "table":
        _append_table(document, node.rows)
    elif node.kind in {"code", "mermaid"}:
        paragraph = document.add_paragraph(style="Handbook Code")
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.18)
        run = paragraph.add_run(node.text)
        run.font.name, run.font.size = "Courier New", Pt(8)
        if node.kind == "mermaid":
            paragraph.text = "Diagram (Mermaid source):\n" + node.text
    elif node.kind == "callout":
        paragraph = document.add_paragraph(node.text)
        properties = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F4F6F9")
        properties.append(shading)
        paragraph.paragraph_format.keep_together = True
    elif node.kind in {"pagebreak", "volumebreak"}:
        return


def build_docx(root: Path, output: Path, toc_page_numbers: dict[str, int] | None = None, paths: Iterable[Path] | None = None) -> Path:
    """Build authoritative DOCX with deterministic static contents."""
    root = root.resolve()
    document = Document()
    _configure_styles(document)
    _set_page(document)
    metadata = _read_yaml(root / "metadata.yaml").get("handbook", {})
    title = str(metadata.get("title", "Enterprise Engineering Manual"))
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.font.name, run.font.size, run.font.bold, run.font.color.rgb = "Calibri", Pt(24), True, RGBColor(11, 37, 69)
    document.add_paragraph("Professional operating standards, controls, and reusable engineering procedures.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    document.add_paragraph("Contents", style="Heading 1")
    all_nodes: list[Node] = []
    for relative in list(paths or manifest_paths(root)):
        all_nodes.extend(parse_document((root / relative).read_text(encoding="utf-8")))
    # Use one heading per manifest document, rather than Markdown heading
    # level: legacy chapters intentionally use flat headings for compact prose.
    toc_rows = manifest_headings(root, paths)
    contents = document.add_table(rows=0, cols=2)
    contents.style = "Light Shading Accent 1"
    contents.autofit = False
    for node in toc_rows:
        cells = contents.add_row().cells
        _add_internal_link(cells[0].paragraphs[0], ("  " * (node.level - 1)) + node.text, _anchor(node.text))
        cells[1].text = str((toc_page_numbers or {}).get(node.text, "?"))
        for paragraph in cells[0].paragraphs + cells[1].paragraphs:
            paragraph.paragraph_format.space_after = Pt(1)
            for run in paragraph.runs:
                run.font.size = Pt(8)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_page_break()
    break_pending = False
    bookmark_id = 1
    for node in all_nodes:
        if node.kind in {"pagebreak", "volumebreak"}:
            break_pending = True
            continue
        _append_node(document, node, page_break_before=break_pending, bookmark_id=bookmark_id)
        if node.kind == "heading":
            bookmark_id += 1
        break_pending = False
    output.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = title
    document.core_properties.author = "Engineering Enablement"
    document.save(output)
    return output


def build_pdf(docx: Path, output: Path) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="handbook-soffice-") as profile:
        environment = {**os.environ, "HOME": profile, "UserInstallation": f"file://{profile}/lo"}
        completed = subprocess.run([str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(output.parent), str(docx)], capture_output=True, text=True, env=environment, check=False)
    converted = output.parent / f"{docx.stem}.pdf"
    if completed.returncode != 0 or not converted.is_file():
        raise RuntimeError(completed.stderr.strip() or completed.stdout.strip() or "LibreOffice PDF conversion failed")
    if converted != output:
        converted.replace(output)
    return output


def write_reference_style(output: Path) -> Path:
    """Create the committed visual reference from the publication's style tokens."""
    document = Document()
    _configure_styles(document)
    _set_page(document)
    document.add_paragraph("Enterprise Engineering Manual", style="Heading 1")
    document.add_paragraph("Publication style reference: compact, readable, and operational.")
    document.add_paragraph("Heading level two", style="Heading 2")
    document.add_paragraph("Heading level three", style="Heading 3")
    document.add_paragraph("A body paragraph demonstrates the production prose rhythm and line spacing.")
    document.add_paragraph("A controlled checklist item", style="List Bullet")
    _append_table(document, [["Field", "Expected treatment"], ["Control", "Stable ID, owner, frequency, and evidence"], ["KPI", "Calculation, target, warning, and anti-gaming guard"]])
    _append_node(document, Node("code", "rtk pytest engineering-handbook/automation/tests -q", language="shell"))
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def _heading_pages(pdf: Path, headings: Iterable[str]) -> dict[str, int]:
    pages = PdfReader(str(pdf)).pages
    result: dict[str, int] = {}
    for page_number, page in enumerate(pages, 1):
        content = page.extract_text() or ""
        for heading in headings:
            if heading in content:
                # The contents page lists every heading before the body does.
                # Retaining the last match selects the actual body location for
                # the unique manifest-document navigation entries.
                result[heading] = page_number
    return result


def _render_pdf(pdf: Path, destination: Path) -> list[Path]:
    destination.mkdir(parents=True, exist_ok=True)
    # Rebuilds may have fewer pages than the preceding publication. Remove only
    # this renderer's deterministic page artifacts before counting fresh output.
    for image in destination.glob("page-*.png"):
        image.unlink()
    prefix = destination / "page"
    completed = subprocess.run([PDFTOPPM, "-png", "-r", "110", str(pdf), str(prefix)], capture_output=True, text=True, check=False)
    pages = sorted(destination.glob("page-*.png"))
    if completed.returncode != 0 or not pages:
        raise RuntimeError(completed.stderr.strip() or "Poppler PNG rendering failed")
    return pages


def mark_visual_review(ledger: Path, reviewer: str) -> Path:
    """Record an explicit human review after the rendered pages have been inspected."""
    record = json.loads(ledger.read_text(encoding="utf-8"))
    if record.get("status") != "rendered":
        raise ValueError("Only a rendered visual-QA ledger may be marked reviewed")
    record["status"] = "passed"
    record["reviewer"] = reviewer
    for page in record["reviewed_pages"]:
        page["status"] = "reviewed"
    ledger.write_text(json.dumps(record, indent=2) + "\n", encoding="utf-8")
    return ledger


def build_publication_spike(root: Path, destination: Path) -> dict[str, Path]:
    """Run the mandatory DOCX-authority publication spike and record page QA."""
    root, destination = root.resolve(), destination.resolve()
    pilot = Path("examples/publication-spike/pilot.md")
    nodes = parse_document((root / pilot).read_text(encoding="utf-8"))
    headings = [node.text for node in nodes if node.kind == "heading"]
    markdown = build_markdown(root, destination / "Enterprise_Engineering_Manual_Pilot.md", [pilot])
    provisional = build_docx(root, destination / "Enterprise_Engineering_Manual_Pilot.provisional.docx", paths=[pilot])
    provisional_pdf = build_pdf(provisional, destination / "Enterprise_Engineering_Manual_Pilot.provisional.pdf")
    pages = _heading_pages(provisional_pdf, headings)
    missing = sorted(set(headings) - set(pages))
    if missing:
        raise RuntimeError(f"Pilot headings missing from provisional PDF: {', '.join(missing)}")
    docx = build_docx(root, destination / "Enterprise_Engineering_Manual_Pilot.docx", pages, [pilot])
    pdf = build_pdf(docx, destination / "Enterprise_Engineering_Manual_Pilot.pdf")
    final_pages = _heading_pages(pdf, headings)
    if any(final_pages.get(item) != page for item, page in pages.items()):
        raise RuntimeError("Static contents page references changed after the final DOCX build")
    pngs = _render_pdf(pdf, destination / "visual-qa/pilot")
    ledger = destination / "visual-qa/pilot-ledger.json"
    ledger.write_text(json.dumps({"status": "rendered", "pages": len(pngs), "reviewed_pages": [{"page": index + 1, "image": image.name, "status": "pending", "findings": [], "correction_cycle": 0} for index, image in enumerate(pngs)], "toc_pages": pages}, indent=2) + "\n", encoding="utf-8")
    return {"markdown": markdown, "docx": docx, "pdf": pdf, "ledger": ledger}


def build_full_publication(root: Path, destination: Path) -> dict[str, Path]:
    """Build the authoritative manual with verified static contents references."""
    root, destination = root.resolve(), destination.resolve()
    headings = [node.text for node in manifest_headings(root)]
    markdown = build_markdown(root, destination / "Enterprise_Engineering_Manual.md")
    provisional = build_docx(root, destination / "Enterprise_Engineering_Manual.provisional.docx")
    provisional_pdf = build_pdf(provisional, destination / "Enterprise_Engineering_Manual.provisional.pdf")
    pages = _heading_pages(provisional_pdf, headings)
    missing = sorted(set(headings) - set(pages))
    if missing:
        raise RuntimeError(f"Manual headings missing from provisional PDF: {', '.join(missing)}")
    docx = destination / "Enterprise_Engineering_Manual.docx"
    pdf = destination / "Enterprise_Engineering_Manual.pdf"
    for _ in range(3):
        build_docx(root, docx, pages)
        build_pdf(docx, pdf)
        final_pages = _heading_pages(pdf, headings)
        if all(final_pages.get(item) == page for item, page in pages.items()):
            break
        pages = final_pages
    else:
        raise RuntimeError("Static contents page references did not stabilize after three manual DOCX builds")
    pngs = _render_pdf(pdf, destination / "visual-qa/final")
    ledger = destination / "visual-qa/final-ledger.json"
    ledger.write_text(json.dumps({"status": "rendered", "pages": len(pngs), "reviewed_pages": [{"page": index + 1, "image": image.name, "status": "pending", "findings": [], "correction_cycle": 0} for index, image in enumerate(pngs)], "toc_pages": pages}, indent=2) + "\n", encoding="utf-8")
    return {"markdown": markdown, "docx": docx, "pdf": pdf, "ledger": ledger}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("root", type=Path)
    parser.add_argument("output_dir", type=Path)
    parser.add_argument("--pilot", action="store_true")
    parser.add_argument("--reference-style", action="store_true")
    parser.add_argument("--approve-pilot", action="store_true")
    parser.add_argument("--reviewer")
    args = parser.parse_args()
    if args.approve_pilot:
        mark_visual_review(args.output_dir / "visual-qa/pilot-ledger.json", args.reviewer or "unspecified")
    elif args.reference_style:
        write_reference_style(args.root / "build/styles/reference.docx")
    elif args.pilot:
        build_publication_spike(args.root, args.output_dir)
    else:
        print(build_full_publication(args.root, args.output_dir)["markdown"])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
